import os
from utility_func import *
from inverted_index import *
from docx import Document
from docx.shared import Cm
from docx.shared import Pt


# establish corpus dictionary
with open("Word_Unigram_Xinhua98Upper.txt", "r") as c:
    corpus = split_corpus(c)

all_doc = []
doc_order = []
for root, dirs, files in os.walk("./SPLIT_DOC_WDID_NEW/", topdown=False):
    for name in files:
        if name[0] == "V":
            doc_order.append(name)
            with open(os.path.join(root, name), "r") as doc:
                f = split_searchfile(doc)
                all_doc.append(f)

all_query = []
query_order = []
for root, dirs, files in os.walk("./QUERY_WDID_NEW_middle", topdown=False):
    for name in files:
        if name[0] == "2":
            query_order.append(name)
            with open(os.path.join(root, name), "r") as q:
                f = split_short_termfile(q)
                f.remove('')
                all_query.append(f)

# establish document database
doc_db = Database()
for doc in doc_order:
    ind = doc_order.index(doc)
    doc_db.add_index(doc, ind)
    doc_db.add_length(ind, len(all_doc[ind]))

# establish document inverted index
inverted_file = InvertedIndex(doc_db)
for text in all_doc:
    inverted_file.index_document(all_doc.index(text), text)

# set up unigram of each document and form a dictionary with docId as a key and dictionary of words probability as a value
# doc_LM = {}
# for doc in doc_order:
#     doc_id = doc_db.get_index(doc)
#     doc_length = doc_db.get_length(doc)
#     wordfreq_dict = inverted_file.lookup_terms_freq(all_doc[doc_id], doc_id)
#     doc_LM[doc] = unigram(wordfreq_dict, doc_length)

# caulculate the query likelihood of documents(old)
# doc_id = doc_db.get_index(doc_order[1])
# doc_length = doc_db.get_length(doc_order[1])
# wordfreq_dict = inverted_file.lookup_terms_freq(all_doc[doc_id], doc_id)
# test_score = modeling(wordfreq_dict, doc_length, corpus, all_query[1])
# print(test_score)

# caulculate the query likelihood of documents
docs_freq = dict()
docs_length = dict()
for doc in doc_order:
    doc_id = doc_db.get_index(doc)
    doc_length = doc_db.get_length(doc)
    wordfreq_dict = inverted_file.lookup_terms_freq(all_doc[doc_id], doc_id)
    docs_freq[doc] = wordfreq_dict
    docs_length[doc] = doc_length

result = []
doc_result = []
for query in all_query:
    q_rank = query_likelihood(docs_freq, docs_length, corpus, query)
    q_order = [s[0] for s in q_rank]
    result.append(q_rank)
    doc_result.append(q_order)


# mAP for query likelihood 
with open('HW1_AssessmentTrainSet.txt', 'r') as relevant_set:
    relevance = split_file(relevant_set, 16)
print(len(relevance))
# compute the mAP
total_ap = 0
for i in range(16):
    mAP = mean_average_precision(doc_result[i], relevance[i])
    total_ap += mAP
total_map = round(total_ap/16, 8)
# print(total_map)
map_answer1 = str(total_map)

# document = Document()

# style = document.styles['Normal']
# font = style.font
# font.name = 'Arial'
# font.size = Pt(12)

# document.add_heading('Information Retrieval HW4', 0)

# p1 = document.add_paragraph('System model: unigram model\n')

# p1.add_run("mean average precision of original query: ")
# p1.add_run(map_answer1)

document.save('90899703Y_HW4.docx')












