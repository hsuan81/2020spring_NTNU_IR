import os 
import math
import numpy as np
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
import function 


def split_termfile(_file) -> list:
    """split the file and store different terms in a list and return the list. """
    answer = []
    for i in _file:
        term = i.split(" ")
        term.remove("-1\n")
        answer.extend(term)
    return answer

def split_searchfile(_file) -> list:
    """split the file and store different terms in a list and return the list. """
    answer = []
    dump = []
    for i in _file:
        term = i.split("\n")
        if len(dump) == 3:
            text = term[0].split(" ")
            text.remove("-1")
            answer.extend(text)
        else:
            dump.append(term)
    return answer

class QueryCorpus:
    def __init__(self, query: list):
        """Count the number of each term in query, build up a term corpus for the query file and return the dictionary"""
        self.query = query
        # compute occurence of each word in query
        query_corpus = {}
        for i in self.query:
            if i not in query_corpus.keys():
                query_corpus[i] = 1
            else:
                query_corpus[i] += 1
        self.corpus = query_corpus
        self.term_indoc = [0] * len(list(self.corpus.keys()))
        
    def item_number(self):
        return len(list(self.corpus.keys()))

    def number_doc(self, document: list):
        """Calculate the number of document containing the word in the query and return a list."""
       
        for i in range(len(document)):
            if document[i] > 0:
                self.term_indoc[i] += 1
        
    
    def length(self, weight_list):
        sum = 0
        for i in weight_list:
            sum += i ** 2
        return math.sqrt(sum)

    def idf(self, number_of_doc: int):
        q_idf = []
        for j in self.term_indoc:
            if j == 0:
                idf = 0
            else:
                idf = math.log((number_of_doc)/(j), 2)  #avoid divided by zero
            q_idf.append(idf)
        return q_idf

    def bm_idf(self, number_of_doc: int):
        q_idf = []
        for i in self.term_indoc:
            idf = math.log((number_of_doc - i + 0.5)/(i + 0.5), 2)
            q_idf.append(idf)
        return q_idf
        

    def weight(self, idf_list) -> list: 
        q_tf = []
        q_idf = idf_list  #
        q_weight = []
        for i in self.corpus.values():
            #f = i
            f = math.log(i, 2) + 1
            q_tf.append(f)
        
        for k in range(self.item_number()): 
            w = q_tf[k] * q_idf[k]
            q_weight.append(w)
        return q_weight

class DocVector:
    """Compute the occurence of query in one doc, and compute the weight, length of doc and bm25 score"""
    def __init__(self, _file: list, query: list):
        frequency = [0] * len(query)
        for i in range(len(query)):
            if query[i] in _file:
                num = _file.count(query[i])
                frequency[i] = num
        self.frequency = frequency
        self.word_length = len(_file)
        
        
    def weight(self, idf: list) -> list: 
        d_idf = idf
        d_tf = []
        d_weight = []
        # compute tf of terms in the doc
        for i in self.frequency:
            if i == 0:
                f = 0
            else:
                #f = i
                f = math.log(i, 2) + 1
            d_tf.append(f)

        for k in range(len(d_idf)):
            w = d_tf[k] * d_idf[k]
            d_weight.append(w)
        return d_weight
        
    def length(self, weight_list):
        sum = 0
        for i in weight_list:
            sum += i ** 2
        return math.sqrt(sum)
    
    def bm_score(self, idf, avg_length, k1, b):
        sum = 0
        for i in range(len(self.frequency)):
            score = idf[i]*(self.frequency[i]*(k1 + 1)/(self.frequency[i]+k1*(1-b+b*self.word_length/avg_length)))
            sum += score
        return sum

def similarity(query_weight, query_length, doc_weight, doc_weight_length):

    product = np.dot(query_weight, doc_weight)
    if doc_weight_length == 0:
        rank = 0
    else:
        rank = product / (doc_weight_length * query_length)
    return rank

def vector_space_model(query_content, doc_order, all_doc):
    que1 = QueryCorpus(query_content)
    all_doc_vector = []
    for passage in all_doc:
        doc1 = DocVector(passage, list(que1.corpus.keys()))
        term_num_doc = que1.number_doc(doc1.frequency)
        all_doc_vector.append(doc1)
    
    
    que1_idf = que1.idf(2265)
    que1_weight = np.array(que1.weight(que1_idf))
    que1_length = que1.length(que1_weight)
    all_doc_weight = []
    all_doc_rank = []
    for i in range(len(all_doc_vector)):
        w = np.array(all_doc_vector[i].weight(que1_idf))
        all_doc_weight.append(w)
        rank = similarity(que1_weight, que1_length, w, all_doc_vector[i].length(w))
        all_doc_rank.append([doc_order[i], rank])
    all_doc_rank.sort(key= lambda x: x[1], reverse=True)
    q1_result = []
    for i in all_doc_rank:
        q1_result.append(i[0])
    return all_doc_rank, q1_result

def BM_25(query_content, doc_order, all_doc):
    k1 = 2.0
    b = 0.75
    que1 = QueryCorpus(query_content)
    all_doc_vector = []
    avg_length = 0
    for passage in all_doc:
        doc1 = DocVector(passage, list(que1.corpus.keys()))
        term_num_doc = que1.number_doc(doc1.frequency)
        all_doc_vector.append(doc1)
        avg_length += doc1.word_length
    avg_length = avg_length / len(all_doc_vector)

    que1_idf = que1.bm_idf(2265)
    
    all_doc_rank = []
    for i in range(len(all_doc_vector)):
        score = all_doc_vector[i].bm_score(que1_idf, avg_length, k1, b)
        all_doc_rank.append([doc_order[i], score])
    all_doc_rank.sort(key= lambda x: x[1], reverse=True)

    q1_result = []
    for i in all_doc_rank:
        q1_result.append(i[0])
    return all_doc_rank, q1_result
    


all_doc = []
doc_order = []
for root, dirs, files in os.walk("./SPLIT_DOC_WDID_NEW", topdown=False):
    for name in files:
        if name[0] == "V":
            doc_order.append(name)
            with open(os.path.join(root, name), "r") as doc:
                f = split_searchfile(doc)
                all_doc.append(f)



all_query = []
query_order = []
for root, dirs, files in os.walk("./QUERY_WDID_NEW", topdown=False):
    for name in files:
        if name[0] == "2":
            query_order.append(name)
            with open(os.path.join(root, name), "r") as q:
                f = split_termfile(q)
                all_query.append(f)


all_result1 = []
for i in range(len(all_query)):
    doc_rank, result = vector_space_model(all_query[i], doc_order, all_doc)
    all_result1.append(result)

all_result2 = []
for i in range(len(all_query)):
    doc_rank, result = BM_25(all_query[i], doc_order, all_doc)
    all_result2.append(result)

with open("VSM_rank_result.txt", "w") as doc:
    for i in range(len(all_result1)):
        for j in range(2265):
            if j == 0:
               doc.write("Query " + str(i+1) + "result\n")
            doc.write(all_result1[i][j] + "\n") 

with open("BM25_rank_result.txt", "w") as doc:
    for i in range(len(all_result2)):
        for j in range(2265):
            if j == 0:
               doc.write("Query " + str(i+1) + "result\n")
            doc.write(all_result2[i][j] + "\n") 

# Compute mean average precision for the search system
with open('HW1_AssessmentTrainSet.txt', 'r') as relevant_set:
    relevance = function.split_file(relevant_set)

# vector space model
total_ap1 = 0
for i in range(16):
    mAP = function.mean_average_precision(all_result1[i], relevance[i])
    total_ap1 += mAP
total_map1 = round(total_ap1/16, 8)
map_answer1 = "mAP = " + str(total_map1)
print(map_answer1)

# BM25 model
total_ap2 = 0
for i in range(16):
    mAP = function.mean_average_precision(all_result2[i], relevance[i])
    total_ap2 += mAP
total_map2 = round(total_ap2/16, 8)
map_answer2 = "mAP = " + str(total_map2)

# document = Document()


# style = document.styles['Normal']
# font = style.font
# font.name = 'Arial'
# font.size = Pt(12)

# document.add_heading('Information Retrieval HW2', 0)

# p1 = document.add_paragraph('System model: Vector space model\n', style = 'List Number')

# p1.add_run("mean average precision: ")
# p1.add_run(map_answer1)

# p2 = document.add_paragraph('System model: BM25\n', style = 'List Number')
# p2.add_run('mean average precision: ')
# p2.add_run(map_answer2)

# document.save('90899703Y_HW2.docx')




        
