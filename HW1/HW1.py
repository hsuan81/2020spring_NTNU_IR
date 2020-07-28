import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm
import math

def split_file(file):
    """split the file by different queries into seperate list element and return one list as a whole. """
    answer = [[]]
    j = 0
    for i in file:
        if i == "\n":
            j += 1
            if j == 16:  # answer list only needs to add 15 additional empty sublists
                pass
            else:
                answer.append([])
        elif i[:3] == "VOM":
            answer[j].append(i)
    return answer

def calculate(retrieve, order, relevant_number):
    """Return recall and precision of each found relevant element."""
    recall = round(retrieve / relevant_number, 4)
    precision = round(retrieve / order, 4)
    return recall, precision

def recall_interval(recall):
    """Return the interval of recall classified to 11 equal intervals of 10 (the field range is 0-100)"""
    return ((recall*10)//1)*10

def compare_interpolate(recall_area):
    """
    o the interpolation within each interval and beyond accoording to the function.
    Return the lists of precisions and interpolated recalls.
    """
    final_r = []
    final_p = []
    for i in range(100, -1, -10):
        final_r.append(i)
        if recall_area[i] != []:
            # interpolate if the max precision is smaller than the larger interval
            if i != 100 and recall_area[i][0][0]*100 < final_p[-1]:
                final_p.append(final_p[-1])
            else:
                final_p.append(recall_area[i][0][0]*100)
        # if no precision is in the interval, use the precision of larger interval
        else:
            final_p.append(final_p[-1])
    return final_p, final_r

def mean_average_precision(answer_set, relevant_set):
    """calculte mean average precision by summing up all precision and 
    divide the sum by teh number of relevant elements."""
    order = 0
    retrieve = 0
    sum = 0
    relevant_number = len(relevant_set)
    for i in range(len(answer_set)):
        order += 1
        for j in relevant_set:
            if answer_set[i][:21] == j[:21]: 
                retrieve += 1
                recall, precision = calculate(retrieve, order, relevant_number)
                # r.append(recall)
                # p.append(precision)
                sum += precision
                if retrieve > len(relevant_set):
                    break
    # compute the mean average precision
    mean_ap = sum/relevant_number
    return mean_ap



def interpolate(answer_set, relevant_set):
    order = 0
    retrieve = 0
    recall_area = {0:[], 10:[], 20:[], 30:[], 40:[], 50:[], 60:[], 70:[], 80:[], 90:[], 100:[]}
    r = []
    p = []
    relevant_number = len(relevant_set)
    for i in range(len(answer_set)):
        order += 1
        for j in relevant_set:
            if answer_set[i][:21] == j[:21]: 
                retrieve += 1
                recall, precision = calculate(retrieve, order, relevant_number)
                recall_area[recall_interval(recall)].append((precision, recall))
                r.append(recall)
                p.append(precision)
                if retrieve > len(relevant_set):
                    break
    
    # interpolation of the precision
    for i in recall_area.values():
        i.sort(reverse = True)
    
    final_p, final_r = compare_interpolate(recall_area)
    final_r = []
    final_p = []
    for i in range(100, -1, -10):
        final_r.append(i)
        if recall_area[i] != []:
            if i != 100 and recall_area[i][0][0]*100 < final_p[-1]:
                # interpolate if the max precision is smaller than the larger interval
                final_p.append(final_p[-1])
            else:
                final_p.append(recall_area[i][0][0]*100)
        else:
            # if no precision is in the interval, use the precision of larger interval
            final_p.append(final_p[-1])

    return final_r, final_p



with open('HW1_ResultsTrainSet.txt', 'r') as answer_set:
    answer = split_file(answer_set)

with open('HW1_AssessmentTrainSet.txt', 'r') as relevant_set:
    relevance = split_file(relevant_set)

total_precision = {x:[] for x in range(100, -1, -10)}
for i in range(16):
    r, p = interpolate(answer[i], relevance[i])
    for i in r:
        total_precision[i].append(p[r.index(i)])


final_precision = []
final_recall = [x for x in range(100, -1, -10)]
for i in total_precision.values():
    sum = 0
    for j in i:
        sum += j
    result = sum/16
    if final_precision != [] and result < final_precision[-1]:
        # interpolate if the max precision is smaller than the larger interval
        final_precision.append(final_precision[-1])
    else:
        final_precision.append(result)

plt.plot(final_recall, final_precision, marker = ".") 
plt.xlabel("Recall")
plt.ylabel("Precision")
plt.title("Interpolated Recall-Precision Curve", loc = 'center')
#set x, y axis to fixed range
plt.axis([0,100,0,100])
plt.savefig("interpolate.png")
plt.close("interpolate.png")
plt.clf()
plt.cla()


# calculate map
total_ap = 0
for i in range(16):
    mAP = mean_average_precision(answer[i], relevance[i])
    total_ap += mAP
total_map = round(total_ap/16, 8)
map_answer = "mAP = " + str(total_map)
print(map_answer)


# calculate dcg
# score range is 0-3 and equally distributed among the relevant set
def assign_score(relevant_set):
    """Assign score to each relevant element in descending order and return the score list."""
    section = len(relevance[0])//3
    score = []
    s = 3
    for i in range(3):
        if s == 1:
            num = len(relevance[0]) - len(score)
            score.extend([s]*num)
        else:
            score.extend([s]*section)
            s -= 1
    return score

def gain(answer_set, relevant_set, score_list):
    """Form a score list based on the answer set and return the rank list."""
    rank_list = []
    order_list = []
    for i in range(len(answer_set)):
        item = answer_set[i][:21] + "\n"
        if item in relevant_set:
            rank_list.append(score_list[relevant_set.index(item)])
            order_list.append(item)
        else:
            rank_list.append(0)
            order_list.append("no")
    c = rank_list.count(0)
    return rank_list

def cumulative_gain(rank_list):
    """Calculate the cumulative gain based on the rank list and return a list."""
    cumulative_set = []
    cumulative_set.append(rank_list[0])
    for i in range(1, len(rank_list)):
        cg = cumulative_set[i-1] + rank_list[i]
        cumulative_set.append(cg)
    return cumulative_set

def discounted_cumulative_gain(rank_list):
    """Calculate the discounted cumulative gain based on the input rank list and return a list."""
    discounted_cg = []
    discounted_cg.append(rank_list[0])
    for i in range(1, len(rank_list)):
        d = rank_list[i]/math.log2(i+1)
        dcg = d + discounted_cg[i-1]
        discounted_cg.append(dcg)
    return discounted_cg

def ideal_dcg(score, answer_set_number):
    """Calculate the ideal discounted cumulative gain based on a descending rank list and return a list."""
    ideal_set = score
    added = answer_set_number - len(score)
    ideal_set.extend([0]*added)
    idgc = discounted_cumulative_gain(ideal_set)
    return idgc

def normalized_dcg(query_number, answer_set, relevant_set, score, rank_list):
    """Calculate normalized discounted cumulative gain of various queries and return a list."""
    total_dcg = []
    total_idcg = []
    for i in range(query_number):
        dcg = discounted_cumulative_gain(rank_list)
        total_dcg.append(dcg)
        idcg = ideal_dcg(score, len(answer_set[i]))
        total_idcg.append(idcg)
    
    final_idcg = 0
    final_dcg = 0
    total_ndcg = []
    for i in range(len(answer_set[0])):
        for j in range(query_number):
            final_idcg += total_idcg[j][i]
            final_dcg += total_dcg[j][i]
        ndcg = final_dcg / final_idcg
        total_ndcg.append(ndcg)  
    return total_ndcg

score = assign_score(relevance[0])
rank = gain(answer[0], relevance[0], score)
cg = cumulative_gain(rank)
discounted_cumulative_gain(rank)
ndcg = normalized_dcg(16, answer, relevance, score, rank)


plt.plot(ndcg, 'g')
plt.xlabel("Answer Set")
plt.title("Normalized Discounted Cumulated Gain", loc = 'center')
plt.axis([0, 2500, 0 , 1])
plt.savefig("NDCG.png")

# combine graph and answer into one document
document = Document()

document.add_heading('Information Retrieval HW1', 0)

p1 = document.add_paragraph('Interpolated precision recall curve', style = 'List Number')
document.add_picture('interpolate.png', width=Cm(12))

p2 = document.add_paragraph('Mean average precision\n', style = 'List Number')
p2.add_run(map_answer)

p3 = document.add_paragraph('Normalized discounted cumulated gain', style = 'List Number')
document.add_picture('NDCG.png', width=Cm(12))

document.save('90899703Y_HW1.docx')




    
                




    


