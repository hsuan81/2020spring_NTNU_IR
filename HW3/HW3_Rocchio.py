import os
from search_system import *
from system_evaluation import *
from function import *
from Rocchio import *
from docx import Document
from docx.shared import Cm
from docx.shared import Pt


all_doc = []
doc_order = []
for root, dirs, files in os.walk("./SPLIT_DOC_WDID_NEW/", topdown=False):
    for name in files:
        if name[0] == "V":
            doc_order.append(name)
            with open(os.path.join(root, name), "r") as doc:
                f = split_searchfile(doc)
                all_doc.append(f)

all_query = []
query_order = []
for root, dirs, files in os.walk("./QUERY_WDID_NEW_middle", topdown=False):
    for name in files:
        if name[0] == "2":
            query_order.append(name)
            with open(os.path.join(root, name), "r") as q:
                f = split_short_termfile(q)
                f.remove('')
                all_query.append(f)
                
query_number = len(all_query)
word_lis = word_collection(all_doc)

# vector space model
all_result1 = []
for i in range(len(all_query)):
    result, ranks = vsm(all_query[i], doc_order, all_doc, word_lis)
    all_result1.append(result)

# create a file to store the result of short query search
with open("search_result_from_short_query", "w") as out_file1:
    for i in range(len(all_result1)):
        out_file1.write("Query " + str(i+1) + "\n")
        for doc in all_result1[i]:
            out_file1.write(doc + "\n")

# mAP for vector space model
with open('HW1_AssessmentTrainSet.txt', 'r') as relevant_set:
    relevance = function.split_file(relevant_set)
# compute the mAP
total_ap1 = 0
for i in range(query_number):
    mAP = function.mean_average_precision(all_result1[i], relevance[i])
    total_ap1 += mAP
total_map1 = round(total_ap1/query_number, 8)
map_answer1 = str(total_map1)


# Rocchio method
# with open('./HW1_AssessmentTrainSet.txt', 'r') as relevant_set:
#     relevance = function.split_file(relevant_set)
remove_newline_symbol(relevance)

docs = WordVector(doc_order, all_doc, word_lis)
docs_freq = docs.docs_frequency()
optimal_q_vector = []
optimal_qlis = []
optimal_q_length = []
num_relevant = 3
for i in range(16):
    q1 = QueryVector(word_lis, all_query[0])
    q1_freq = q1.query_frequency()
    q1_idf = q1.idf(docs_freq)
    q1_tf_idf = q1.query_tf_idf(q1_idf)
    q_expansion = RocchioQuery(q1_tf_idf, all_result1[i], q1_idf, word_lis)
    relevant_texts = relevant_text(doc_order, all_doc, relevance[i])
    new_q_vector, new_q, length = q_expansion.rocchio(relevant_texts, relevance[i], num_relevant)
    optimal_q_vector.append(new_q_vector)
    optimal_qlis.append(new_q)
    optimal_q_length.append(length)
    #print(i, optimal_q_vector[i])


# vector space model with new query
new_result = []
for i in range(len(optimal_q_vector)):
    q = QueryVector(word_lis, optimal_qlis[i])
    docs = WordVector(doc_order, all_doc, word_lis)
    q_idf = q.idf(docs.docs_frequency())
    docs_tf_idf = docs.docs_tf_idf(q_idf)
    docs_len = docs.length(docs_tf_idf)
    ranks = docs_similarity(optimal_q_vector[i], optimal_q_length[i], docs_tf_idf, docs_len)
    ranks.sort(key=lambda x: x[1], reverse=True)
    result = []
    for i in range(len(ranks)):
        result.append(ranks[i][0])
    new_result.append(result)


# mAP after query expansion
total_ap2 = 0
for i in range(query_number):
    mAP2 = function.mean_average_precision(new_result[i], relevance[i])
    total_ap2 += mAP2
total_map2 = round(total_ap2/query_number, 8)
map_answer2 = str(total_map2)
#print("after", map_answer2)

# create a file to store the sorting after query expansion
with open("search_result_after_query_expansion", "w") as out_file2:
    for i in range(len(new_result)):
        out_file2.write("Query " + str(i+1) + "\n")
        for doc in new_result[i]:
            out_file2.write(doc + "\n")


# document = Document()

# style = document.styles['Normal']
# font = style.font
# font.name = 'Arial'
# font.size = Pt(12)

# document.add_heading('Information Retrieval HW3', 0)

# p1 = document.add_paragraph('System model: Vector space model\n', style = 'List Number')

# p1.add_run("mean average precision of original query: ")
# p1.add_run(map_answer1)

# p2 = document.add_paragraph('Query expansion method: Rocchio method \n', style = 'List Number')
# p2.add_run('number of relevant documents used in this round: ' + str(num_relevant) + "\n")
# p2.add_run('mean average precision after query expansion: ')
# p2.add_run(map_answer2)

# document.save('90899703Y_HW3.docx')







